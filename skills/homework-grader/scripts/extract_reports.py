"""
Batch extract .doc/.docx student reports to text files.
Handles OLE2 .doc format via Word COM, .docx via python-docx.
"""
import os, sys, subprocess, re

def extract_docx(docx_path, output_path, label):
    """Extract text from .docx using python-docx."""
    try:
        from docx import Document
        doc = Document(docx_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for tbl in doc.tables:
            for row in tbl.rows:
                rt = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                if rt: parts.append(rt)
        text = "\n".join(parts)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"{label}\n{'='*60}\n{text}")
        return len(text)
    except Exception as e:
        print(f"  FAIL (.docx): {e}")
        return 0

def extract_doc_com(doc_path, output_path, label):
    """Extract text from .doc (OLE2) using Word COM via PowerShell."""
    ps_cmd = f"""
$w = New-Object -ComObject Word.Application
$w.Visible = $false
try {{
    $d = $w.Documents.Open('{doc_path}')
    $t = $d.Content.Text
    [System.IO.File]::WriteAllText('{output_path}', '{label}' + $t, [System.Text.Encoding]::UTF8)
    $d.Close()
    Write-Output $t.Length
}} catch {{
    Write-Output 'ERROR'
}} finally {{
    $w.Quit()
}}
"""
    try:
        result = subprocess.run(['powershell', '-Command', ps_cmd], 
                              capture_output=True, text=True, timeout=60)
        output = result.stdout.strip()
        if output and output != 'ERROR' and output.isdigit():
            return int(output)
        return 0
    except:
        return 0

def batch_extract(base_dir, output_dir, student_list):
    """
    Extract all student reports.
    
    student_list: list of tuples (student_id, student_name, report_rel_path)
    """
    os.makedirs(output_dir, exist_ok=True)
    
    for sid, sname, rel_path in student_list:
        full_path = os.path.join(base_dir, rel_path)
        label = f"Student: {sid}-{sname}\nFile: {os.path.basename(full_path)}"
        out_path = os.path.join(output_dir, f"{sid}-{sname}.txt")
        
        if os.path.exists(out_path) and os.path.getsize(out_path) > 500:
            print(f"SKIP: {sid}-{sname} (already extracted)")
            continue
        
        if not os.path.exists(full_path):
            print(f"MISS: {sid}-{sname} (file not found)")
            continue
        
        ext = os.path.splitext(full_path)[1].lower()
        if ext == '.docx':
            size = extract_docx(full_path, out_path, label)
        else:
            size = extract_doc_com(full_path, out_path, label)
        
        if size > 500:
            print(f"OK: {sid}-{sname}: {size} chars")
        else:
            print(f"FAIL: {sid}-{sname}")

if __name__ == '__main__':
    import json
    if len(sys.argv) < 3:
        print("Usage: python extract_reports.py <base_dir> <output_dir> [student_list.json]")
        sys.exit(1)
    
    base_dir = sys.argv[1]
    output_dir = sys.argv[2]
    
    if len(sys.argv) > 3:
        with open(sys.argv[3], 'r', encoding='utf-8') as f:
            student_list = json.load(f)
    else:
        student_list = []
        for entry in os.listdir(base_dir):
            path = os.path.join(base_dir, entry)
            if os.path.isdir(path):
                parts = entry.split('-', 1)
                sid = parts[0]
                sname = parts[1] if len(parts) > 1 else ''
                for f in os.listdir(path):
                    if f.endswith(('.doc', '.docx')) and '报告' in f:
                        student_list.append((sid, sname, os.path.join(entry, f)))
    
    batch_extract(base_dir, output_dir, student_list)
